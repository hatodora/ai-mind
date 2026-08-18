#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
思索 / Mindmap 取扱説明書 — Word 版の組版。

PDF 版（manual.html）と同じ構成・同じ図版を使い、11ページに収める。
2段組みは «罫線のない表» で作っている。Word で段組みを使うと
ページをまたいだときに流れてしまうため。

書体だけは差し替えてある（PDF の見出しは Satoshi、
こちらはどの環境にもある Arial ＋ 游ゴシック）。

    python3 build_docx.py
"""

import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

HERE = os.path.dirname(os.path.abspath(__file__))


def asset(name):
    return os.path.join(HERE, "assets", name)


# ---------------------------------------------------------------- 色と書体

INK = RGBColor(0x24, 0x26, 0x1F)
MUTED = RGBColor(0x6B, 0x6F, 0x63)
FAINT = RGBColor(0x9B, 0xA1, 0x93)
ACCENT = RGBColor(0x0A, 0xBA, 0xB5)
ACCENT_SOFT = RGBColor(0x0A, 0x8F, 0x8B)
ACCENT_DEEP = RGBColor(0x06, 0x6E, 0x6A)
WARM_INK = RGBColor(0x7D, 0x53, 0x10)
DANGER = RGBColor(0xD4, 0x3D, 0x3D)

PANEL = "F6F5F2"
TINT = "E6F7F6"
TINT_WARM = "FAF2E3"
LINE = "E0DFD9"
HAIR = "EEEDE8"
DARK = "24261F"

LATIN = "Arial"
CJK = "Yu Gothic"

CENTER = WD_ALIGN_PARAGRAPH.CENTER
RIGHT = WD_ALIGN_PARAGRAPH.RIGHT


# ---------------------------------------------------------------- XML の小道具


def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    return e


def style_run(run, size=9.5, bold=False, color=INK, spacing=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.font.name = LATIN
    rPr = run._element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = _el("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:ascii"), LATIN)
    rf.set(qn("w:hAnsi"), LATIN)
    rf.set(qn("w:eastAsia"), CJK)
    if spacing is not None:  # 字間。1/20 pt 単位
        rPr.append(_el("w:spacing", val=int(spacing * 20)))
    return run


def shade_cell(cell, fill):
    cell._tc.get_or_add_tcPr().append(_el("w:shd", val="clear", color="auto", fill=fill))


def cell_borders(cell, color=LINE, sz=6, sides=("top", "left", "bottom", "right")):
    borders = _el("w:tcBorders")
    for s in ("top", "left", "bottom", "right"):
        borders.append(
            _el("w:" + s, val="single" if s in sides else "nil", sz=sz, space="0",
                color=color)
        )
    cell._tc.get_or_add_tcPr().append(borders)


def set_cell_margins(cell, top=2.5, left=3.5, bottom=2.5, right=3.5):
    mar = _el("w:tcMar")
    for tag, mm in (("top", top), ("start", left), ("bottom", bottom), ("end", right)):
        mar.append(_el("w:" + tag, w=int(mm * 56.7), type="dxa"))
    cell._tc.get_or_add_tcPr().append(mar)


def no_borders(table):
    borders = _el("w:tblBorders")
    for s in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_el("w:" + s, val="nil"))
    table._tbl.tblPr.append(borders)


# ---------------------------------------------------------------- 段落

# 次に作る «本文の» 段落を新しいページから始めたい、という予約。
# 空段落＋改ページだと白紙のページが生まれるので、この形にしている。
_break_pending = False


def new_page():
    global _break_pending
    _break_pending = True


def _consume_break(paragraph, container):
    global _break_pending
    if _break_pending and container is doc:
        paragraph.paragraph_format.page_break_before = True
        _break_pending = False


def line(container, runs, before=0, after=0, spacing=1.9, align=None, keep=False):
    """runs = [(文字列, 大きさ, 太字, 色, 字間), ...]。後ろは省略可。"""
    p = container.add_paragraph()
    _consume_break(p, container)
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = spacing
    if align is not None:
        p.alignment = align
    if keep:
        pf.keep_with_next = True
    for r in runs:
        style_run(p.add_run(r[0]),
                  r[1] if len(r) > 1 else 9.5,
                  r[2] if len(r) > 2 else False,
                  r[3] if len(r) > 3 else INK,
                  r[4] if len(r) > 4 else None)
    return p


def para(container, text="", size=9.5, bold=False, color=INK, before=0, after=0,
         spacing=1.9, align=None, tracking=None, keep=False):
    return line(container, [(text, size, bold, color, tracking)] if text else [],
                before, after, spacing, align, keep)


def rich(container, parts, size=9.5, before=0, after=0, spacing=1.9, base=INK):
    """[(文字列, 太字?), ...]。太字は濃く、地は muted に落とす。"""
    runs = []
    for it in parts:
        bold = it[1] if len(it) > 1 else False
        runs.append((it[0], size, bold, INK if bold else base))
    return line(container, runs, before, after, spacing)


def rule(container, color=LINE, before=6, after=6):
    p = container.add_paragraph()
    _consume_break(p, container)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    pBdr = _el("w:pBdr")
    pBdr.append(_el("w:bottom", val="single", sz=6, space="0", color=color))
    p._p.get_or_add_pPr().append(pBdr)
    return p


def spacer(container, mm=6):
    para(container, "", size=1, after=mm * 2.83, spacing=1.0)


def picture(container, path, width_mm, align=None, before=0, after=0):
    p = container.add_paragraph()
    _consume_break(p, container)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(path, width=Mm(width_mm))
    return p


# ---------------------------------------------------------------- 表

def add_table(container, rows, cols, widths_mm):
    if hasattr(container, "add_table") and container is not doc:
        t = container.add_table(rows, cols)          # セルの中
    else:
        t = doc.add_table(rows=rows, cols=cols)      # 本文
    t.autofit = False
    t._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))
    for i, w in enumerate(widths_mm):
        t.columns[i].width = Mm(w)
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Mm(widths_mm[i])
    no_borders(t)
    return t


def fix_cells():
    """
    表のセルは «段落で終わる» のが OOXML の決まり。
    入れ子の表を最後に置いたままだと Word が独自に段落を補い、
    行の高さが読めなくなってページ送りが暴れる。
    保存前にまとめて、高さのない段落を足しておく。
    """
    body = doc.element.body
    for tc in body.iter(qn("w:tc")):
        kids = [c for c in tc if c.tag != qn("w:tcPr")]
        if kids and kids[-1].tag == qn("w:p"):
            continue
        p = OxmlElement("w:p")
        pPr = OxmlElement("w:pPr")
        pPr.append(_el("w:spacing", before="0", after="0", line="20",
                       lineRule="exact"))
        rPr = OxmlElement("w:rPr")
        rPr.append(_el("w:sz", val="2"))
        pPr.append(rPr)
        p.append(pPr)
        tc.append(p)


def blank(cell, margins=(0, 0, 0, 0)):
    """セルの既定の空段落を取り除いて、余白を決める。"""
    p = cell.paragraphs[0]._p
    p.getparent().remove(p)
    set_cell_margins(cell, *margins)
    return cell


# ---------------------------------------------------------------- 意味のある塊

def eyebrow(text):
    para(doc, text, size=7, bold=True, color=ACCENT_SOFT, spacing=1.2, after=4,
         tracking=1.6, keep=True)


def h1(text):
    para(doc, text, size=21, bold=True, spacing=1.3, after=5, keep=True)


def lead(text):
    para(doc, text, size=10, color=MUTED, spacing=1.85, after=13)


def micro(container, text, before=0, after=4):
    # keep_with_next は付けない。直後が表のとき、
    # Word はその表を «次のページごと» 送ってしまい白紙が生まれる
    para(container, text, size=6.8, bold=True, color=FAINT, spacing=1.3,
         before=before, after=after, tracking=1.3)


def tip(container, label, parts, warm=False, width=None):
    base = WARM_INK if warm else ACCENT_DEEP
    t = add_table(container, 1, 1, [width or 170])
    c = blank(t.cell(0, 0), (4, 5, 4, 5))
    shade_cell(c, TINT_WARM if warm else TINT)
    para(c, label, size=6.8, bold=True, color=base, spacing=1.3, after=2, tracking=1.3)
    line(c, [(it[0], 8.2, it[1] if len(it) > 1 else False, base) for it in parts],
         spacing=1.8)
    return t


def notebox(container, mascot, title, paragraphs, width=170, mascot_mm=22):
    t = add_table(container, 1, 2, [mascot_mm + 9, width - mascot_mm - 9])
    left = blank(t.cell(0, 0), (5, 5, 5, 2))
    right = blank(t.cell(0, 1), (5, 2, 5, 5))
    for c in (left, right):
        shade_cell(c, PANEL)
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    picture(left, asset(mascot), mascot_mm)
    para(right, title, size=10, bold=True, spacing=1.5, after=3)
    for i, parts in enumerate(paragraphs):
        rich(right, parts, size=8.6, after=0 if i == len(paragraphs) - 1 else 3,
             spacing=1.9, base=MUTED)
    return t


def step(num, title, fig, fig_mm, body, tip_label=None, tip_parts=None,
         warm=False, reverse=False):
    """図と文を左右に並べる。reverse で図を右に。"""
    FIG, TXT = 84, 82
    t = add_table(doc, 1, 2, [TXT, FIG] if reverse else [FIG, TXT])
    fc = t.cell(0, 1 if reverse else 0)
    bc = t.cell(0, 0 if reverse else 1)
    blank(fc, (0, 4, 0, 0) if reverse else (0, 0, 0, 4))
    blank(bc, (0, 0, 0, 4) if reverse else (0, 4, 0, 0))
    fc.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    picture(fc, asset(fig), fig_mm)

    line(bc, [(num, 16, True, ACCENT)], spacing=1.1, after=3)
    para(bc, title, size=13.5, bold=True, spacing=1.45, after=4)
    for parts in body:
        rich(bc, parts, size=9, after=4, spacing=1.9, base=MUTED)
    if tip_label:
        tip(bc, tip_label, tip_parts, warm=warm, width=TXT - 4)
    return t


def kv_table(header, rows, widths, highlight=None):
    t = add_table(doc, len(rows) + (1 if header else 0), len(widths), widths)
    off = 0
    if header:
        for i, htext in enumerate(header):
            c = blank(t.cell(0, i), (1, 2.5, 2.5, 2.5))
            cell_borders(c, LINE, sides=("bottom",))
            para(c, htext, size=6.8, bold=True, color=FAINT, spacing=1.3, tracking=1.2)
        off = 1
    for j, row in enumerate(rows):
        for i, text in enumerate(row):
            c = blank(t.cell(off + j, i), (3, 2.5, 3, 2.5))
            cell_borders(c, HAIR, sides=("bottom",))
            key = i == 0
            color = INK if key else MUTED
            if highlight == j and key:
                color = ACCENT_SOFT
            para(c, text, size=8.6, bold=key, color=color, spacing=1.8)
    return t


# ================================================================ 文書

doc = Document()

normal = doc.styles["Normal"]
normal.font.name = LATIN
normal.font.size = Pt(9.5)
normal.font.color.rgb = INK
normal.element.rPr.rFonts.set(qn("w:eastAsia"), CJK)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.9

sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(20)
sec.top_margin, sec.bottom_margin = Mm(19), Mm(16)
sec.different_first_page_header_footer = True

foot = sec.footer.paragraphs[0]
foot.paragraph_format.line_spacing = 1.2
style_run(foot.add_run("思索 / Mindmap — はじめかた"), 6.8, False, FAINT, spacing=0.8)
foot.add_run("\t\t")
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
_r = OxmlElement("w:r")
_rPr = OxmlElement("w:rPr")
for tag, val in (("sz", "14"), ("color", "6B6F63")):
    e = OxmlElement("w:" + tag)
    e.set(qn("w:val"), val)
    _rPr.append(e)
_rPr.append(OxmlElement("w:b"))
_r.append(_rPr)
fld.append(_r)
foot._p.append(fld)


# ---------------------------------------------------------------- 1. 表紙

para(doc, "思索 / Mindmap", size=12, bold=True, spacing=1.4)
para(doc, "はじめかた", size=40, bold=True, spacing=1.15, before=62, after=6)
para(doc, "QUICK START GUIDE", size=8, bold=True, color=ACCENT_SOFT, spacing=1.4,
     after=20, tracking=3.2)
para(doc, "人間の脳で、考えよう。", size=11, color=MUTED, spacing=2.0)
para(doc, "行き詰まったら AI と対話し、想像を膨らませる。", size=11, color=MUTED,
     spacing=2.0)
picture(doc, asset("m-wave.png"), 54, align=RIGHT, before=52)
rule(doc, before=16, after=4)
para(doc, "取扱説明書 — 第 1 版", size=7.4, color=MUTED, spacing=1.3, tracking=0.6)


# ---------------------------------------------------------------- 2. はじめる前に

new_page()
eyebrow("01 — BEFORE YOU START")
h1("はじめる前に")
lead("用意するものは、これだけです。\nアプリのインストールも、道具の組み立ても要りません。")

PARTS = [
    ("ico-phone.png", "スマートフォン", "パソコン・タブレットでも", "×1", False),
    ("ico-browser.png", "ブラウザ", "Safari / Chrome など", "×1", False),
    ("ico-wifi.png", "インターネット", "AI に聞くときだけ必要", "×1", False),
    ("ico-clock.png", "15分ほどの時間", "最初の1枚が書けます", "×1", False),
    ("ico-mail.png", "メールアドレス", "保存を残したいときだけ", "任意", True),
    ("ico-none.png", "アプリのインストール", "要りません", "0", True),
]
grid = add_table(doc, 2, 3, [56.6, 56.6, 56.6])
for i, (ico, name, desc, qty, opt) in enumerate(PARTS):
    c = blank(grid.cell(i // 3, i % 3), (4.5, 4, 5, 4))
    cell_borders(c, LINE)
    if opt:
        shade_cell(c, PANEL)
    para(c, qty, size=7.6, bold=True, color=FAINT if opt else ACCENT_SOFT,
         spacing=1.2, after=1, align=RIGHT)
    picture(c, asset(ico), 14, align=CENTER, after=3)
    para(c, name, size=9, bold=True, spacing=1.5, align=CENTER)
    para(c, desc, size=7.4, color=MUTED, spacing=1.6, align=CENTER)

spacer(doc, 9)
notebox(doc, "m-sit.png", "アカウントは無くても、すぐ使えます", [
    [("ログインしなくてもマップは作れます。ただしその場合、マップは",),
     ("この端末だけ", True), ("に保存され、最終更新から",), ("30日", True),
     ("で消えます。",)],
    [("ログインすると、クラウドに保存されて",), ("どの端末からでも開ける", True),
     ("ようになります。端末に残っているマップは、あとからまとめて取り込めます。",)],
])


# ---------------------------------------------------------------- 3. 各部の名称

new_page()
eyebrow("02 — THE PARTS")
h1("各部の名称")
lead("マップを書く画面は、大きく3つに分かれています。")

picture(doc, asset("editor-callouts.png"), 162, after=9)

legend = add_table(doc, 2, 2, [85, 85])
LEG = [
    ("A", "パネル",
     "テーマ・AI の設定・言葉の入力。操作はほぼここで行います。"
     "スマートフォンでは画面下から引き出すシートになります。"),
    ("B", "キャンバス",
     "マップ本体。指やマウスでドラッグして動かし、つまんで拡大縮小できます。"
     "ノードを押すと選べます。"),
    ("C", "シコウ",
     "左下にいるキャラクター。さわると反応します。押すとメニューが開きます。"),
]
for i, (tag, name, desc) in enumerate(LEG):
    c = blank(legend.cell(i // 2, i % 2), (2, 0 if i % 2 == 0 else 5, 4,
                                           5 if i % 2 == 0 else 0))
    line(c, [(tag + "   ", 9, True, ACCENT_SOFT), (name, 9, True, INK)],
         spacing=1.5, after=1)
    para(c, desc, size=8.4, color=MUTED, spacing=1.8)

rule(doc, before=8, after=7)
micro(doc, "ノードの3つの状態")

states = add_table(doc, 2, 3, [56.6, 56.6, 56.6])
for i, (img, title, desc) in enumerate([
    ("crop-node-root.png", "テーマ", "中心にある問い。1枚に1つ"),
    ("crop-node-user.png", "自分の言葉", "白いカード"),
    ("crop-node-ai.png", "AI の言葉", "破線＋うすい色"),
]):
    picture(blank(states.cell(0, i)), asset(img), 40, align=CENTER)
    c = blank(states.cell(1, i))
    para(c, title, size=8.4, bold=True, spacing=1.4, before=3, align=CENTER)
    para(c, desc, size=7.4, color=MUTED, spacing=1.6, align=CENTER)


# ---------------------------------------------------------------- 4. パネルの中身

new_page()
eyebrow("02 — THE PARTS")
h1("パネルの中身")
lead("上から順に、こう並んでいます。")

DETAILS = [
    ("1", "ターン表示", "crop-turn.png", 42,
     "いまが「あなたの番」か「AI の番」か。基本はずっとあなたの番です。"),
    ("2", "AI アシスト", "crop-level.png", 74,
     "AI をどれくらい使うかを4段階から選びます。あとからいつでも変えられます。"),
    ("3", "AI ゲージ", "crop-gauge.png", 74,
     "● 1つが、AI に相談できる1回分。自分でノードを足すと溜まります。"),
    ("4", "選択中のノード", "crop-selected.png", 74,
     "いま選んでいる言葉。新しい言葉は、ここにぶら下がります。"),
]
det = add_table(doc, 2, 2, [85, 85])
for i, (n, title, img, w, desc) in enumerate(DETAILS):
    c = blank(det.cell(i // 2, i % 2), (0, 0 if i % 2 == 0 else 5, 9,
                                        5 if i % 2 == 0 else 0))
    line(c, [(n + "   ", 9.4, True, ACCENT_SOFT), (title, 9.4, True, INK)],
         spacing=1.4, after=3)
    picture(c, asset(img), w, after=3)
    para(c, desc, size=8, color=MUTED, spacing=1.85)

spacer(doc, 3)
last = add_table(doc, 1, 2, [85, 85])
lc = blank(last.cell(0, 0), (0, 0, 0, 5))
line(lc, [("5   ", 9.4, True, ACCENT_SOFT), ("入力欄", 9.4, True, INK)],
     spacing=1.4, after=3)
picture(lc, asset("crop-input.png"), 56, after=3)
para(lc, "思いついた言葉を書いて「追加する」。すぐ下の「AI にアイデアを聞く」は、"
         "ゲージが溜まると押せるようになります。",
     size=8, color=MUTED, spacing=1.85)
rc = blank(last.cell(0, 1), (24, 5, 0, 0))
picture(rc, asset("m-read.png"), 17, after=4)
para(rc, "ひとこと", size=9.4, bold=True, spacing=1.4, after=3)
rich(rc, [("ノードは",), ("短い言葉", True),
          ("ほど扱いやすくなります。文章ではなく、単語やひとことで。",)],
     size=8, spacing=1.85, base=MUTED)


# ---------------------------------------------------------------- 5–7. 6つのステップ

new_page()
eyebrow("03 — SIX STEPS")
h1("6つのステップ")
lead("はじめの1枚は、この順でできあがります。")

step("1", "テーマを決める", "new-mobile-top.png", 78,
     [[("いま気になっていることを、",), ("ひとつだけ", True),
       ("書きます。うまい言葉にしなくて大丈夫。"
        "「転職について考えたい」くらいで十分です。",)],
      [("下に並ぶ例を押すと、そのまま入ります。",)]],
     "TIP", [("テーマは1枚に1つ。",), ("迷ったら大きいほうを選ぶ", True),
             ("と、あとから枝を伸ばしやすくなります。",)])

spacer(doc, 11)
step("2", "自分の言葉で、5つ広げる", "canvas-band.png", 78,
     [[("中心のことばを押して選び、思いついた言葉を書いて「追加する」。これを",),
       ("5回", True), ("くり返します。",)],
      [("最初の5個は、AI に頼らず自分で出します。",)]],
     "なぜ", [("先に AI の答えを見てしまうと、",), ("自分の考えが出てこなくなる", True),
              ("から。だから最初の5個は、わざと自分の番にしてあります。",)],
     warm=True, reverse=True)

new_page()
micro(doc, "6つのステップ — つづき", after=0)
spacer(doc, 3)
step("3", "AI が解禁される", "crop-gauge.png", 78,
     [[("5個ぶら下げると、ゲージに ● が点きます。",),
       ("「AI にアイデアを聞く」", True), ("が押せるようになります。",)],
      [("● 1つが、相談1回分です。使うと減り、自分でノードを足すとまた溜まります。",)]],
     "TIP", [("ボタンが押せないときは、ゲージが足りないか、通信が切れています。",)])

spacer(doc, 11)
step("4", "AI に聞いて、自分で選ぶ", "crop-suggest.png", 70,
     [[("AI が2〜3個の案を出します。",),
       ("採用するかどうかを決めるのは、あなたです。", True)],
      [("いる案だけ「＋」で採用。ぜんぶ違うと思ったら「スキップ」で構いません。",)]],
     "注意", [("「全部採用」はゲージを大きく使います。", True),
              ("採用した数 × 3個ぶんのノードを自分で足すまで、次は聞けません。",)],
     warm=True, reverse=True)

new_page()
micro(doc, "6つのステップ — つづき", after=0)
spacer(doc, 3)
step("5", "わからない言葉を、解く", "vignette.png", 80,
     [[("ノードを選んで ",), ("「? このノードがわからない」", True),
       (" を押すと、AI がやさしい言葉で説明します。",)],
      [("説明のむずかしさは、登録した誕生日から自動で調整されます。"
        "小学生には小学生の言葉で返ります。",)]],
     "TIP", [("ノードが散らかってきたら ",), ("「整える」", True),
             ("。ツリー状に並べ直します。",)])

spacer(doc, 11)
step("6", "結論を出して、完成にする", "crop-ring.png", 78,
     [[("「AI に全体をレビューしてもらう」", True),
       ("を押すと、マップ全体を読んで、次の一手を提案します。"
        "根拠にしたノードは光ります。",)],
      [("読んだら ",), ("「マップを一時的に保存する」", True),
       ("。完成の印がつき、バッジがもらえます。",)]],
     "あとで見る", [("ホーム画面の輪は、",), ("自分の考えと AI の提案の割合", True),
                    ("です。自分の割合が高いほど、"
                     "自分の頭で考えた1枚になっています。",)],
     reverse=True)


# ---------------------------------------------------------------- 8. ゲージのしくみ

new_page()
eyebrow("04 — HOW IT WORKS")
h1("AI ゲージのしくみ")
lead("このアプリで AI が出し惜しみをするのは、あなたに先に考えてほしいからです。")

flow = add_table(doc, 1, 5, [50, 12, 46, 12, 50])
FLOW = [("自分でノードを足す", "1個ごとに溜まる", False),
        None,
        ("AI ゲージ", "● ● ○ ○ ○", True),
        None,
        ("AI に相談する", "● 1つ を使う", False)]
for i, item in enumerate(FLOW):
    c = flow.cell(0, i)
    c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if item is None:
        blank(c)
        para(c, "→", size=13, bold=True, color=ACCENT_SOFT, spacing=1.2, align=CENTER)
        para(c, "＋" if i == 1 else "−", size=8, bold=True, color=ACCENT_SOFT,
             spacing=1.2, align=CENTER)
        continue
    title, desc, boxed = item
    blank(c, (4.5, 3.5, 4.5, 3.5))
    cell_borders(c, DARK if boxed else LINE, sz=10 if boxed else 6)
    if not boxed:
        shade_cell(c, PANEL)
    para(c, title, size=9, bold=True, spacing=1.4, align=CENTER)
    para(c, desc, size=7.6, color=ACCENT_SOFT if boxed else MUTED, spacing=1.5,
         align=CENTER)

spacer(doc, 11)
micro(doc, "溜まりかたを選ぶ — AI アシスト")
kv_table(["レベル", "溜まりかた", "向いている人"],
         [["たっぷり", "ノード1個 = 相談2回分", "とにかく手数がほしい"],
          ["標準（既定）", "ノード1個 = 相談1回分", "迷ったらこれ"],
          ["ひかえめ", "ノード3個 = 相談1回分", "自分の頭で粘りたい"],
          ["AI なし", "提案を使わない", "完全に自力で書きたい"]],
         [40, 62, 68], highlight=1)

rule(doc, before=9, after=7)
micro(doc, "3つの決まりごと")
kv_table(None,
         [["はじめの5個",
           "自分で5個ぶら下げるまで、AI はロックされています。"
           "到達すると相談1回分が渡されます。"],
          ["全部採用",
           "まとめて採用すると、採用した数 × 3個ぶんのノードを足すまで次は聞けません。"],
          ["お助け",
           "30個以上に広げ、AI の割合が半分以下で、3分ほど手が止まったとき。"
           "ゲージと関係なく1回だけ相談できます。"]],
         [34, 136])

spacer(doc, 11)
notebox(doc, "m-sleep.png", "手が止まっても、急かされません", [
    [("このアプリは通知やポップアップで先を促しません。",), ("お助け", True),
     ("は、条件がそろったときにパネルの中へ静かに出てくるだけです。",)],
    [("気づかなければ、そのまま自分のペースで書き続けてください。"
      "それがいちばん良い使い方です。",)],
])


# ---------------------------------------------------------------- 9. コツ

new_page()
eyebrow("05 — DO & DON'T")
h1("うまくいくコツ")
lead("同じ機能でも、使い方で仕上がりが変わります。")

DO = [("まず自分の言葉で書く", "きれいでなくていい。思いついた順で、頭にある言葉のまま置きます。"),
      ("ノードは短く", "単語かひとこと。長い文章は枝を伸ばしにくくなります。"),
      ("行き詰まってから AI に聞く", "手が止まった場所でこそ、提案がよく効きます。"),
      ("いらない提案は捨てる", "スキップしてもゲージは減りません。選ぶことも思考のうちです。"),
      ("迷ったら「整える」", "並べ直すと、抜けている枝が見えてきます。")]
DONT = [("最初から AI に全部聞く", "そもそも押せません。自分で5個書くまでロックされています。"),
        ("「全部採用」を連打する", "ゲージが大きく減り、しばらく相談できなくなります。"),
        ("言葉をきれいにしようとする", "整えるのは最後。書く手が止まるほうが損です。"),
        ("1つのノードに文章を詰める", "読み返せなくなり、AI の提案もぼやけます。"),
        ("ログインせずに大事な1枚を書く", "端末にしか残らず、30日で消えます。")]

cols = add_table(doc, 1, 2, [83, 83])
for i, (icon, label, color, items) in enumerate([
    ("ico-check.png", "こうする", ACCENT_DEEP, DO),
    ("ico-cross.png", "こうしない", DANGER, DONT),
]):
    c = blank(cols.cell(0, i), (0, 0 if i == 0 else 5, 0, 5 if i == 0 else 0))
    hd = add_table(c, 1, 2, [10, 62])
    picture(blank(hd.cell(0, 0)), asset(icon), 8)
    hc = blank(hd.cell(0, 1), (0, 2, 0, 0))
    hc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para(hc, label, size=12, bold=True, color=color, spacing=1.3)
    spacer(c, 3)
    for j, (title, desc) in enumerate(items):
        p = para(c, title, size=9, bold=True, spacing=1.5, before=5, after=1)
        if j:
            pBdr = _el("w:pBdr")
            pBdr.append(_el("w:top", val="single", sz=6, space="6", color=HAIR))
            p._p.get_or_add_pPr().append(pBdr)
        para(c, desc, size=7.8, color=MUTED, spacing=1.8, after=3)

rule(doc, before=14, after=9)
closing = add_table(doc, 1, 2, [30, 140])
picture(blank(closing.cell(0, 0), (0, 0, 0, 4)), asset("m-work.png"), 25)
rc = blank(closing.cell(0, 1), (0, 4, 0, 0))
rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
para(rc, "迷ったら、この1行", size=6.8, bold=True, color=FAINT, spacing=1.3, after=3,
     tracking=1.3)
para(rc, "AI に聞く前に、自分の言葉を ひとつ足す。", size=14, bold=True, spacing=1.5,
     after=4)
para(rc, "それだけで、AI の提案は自分の考えに寄っていきます。", size=8.4, color=MUTED,
     spacing=1.8)


# ---------------------------------------------------------------- 10. 困ったときは

new_page()
eyebrow("06 — TROUBLESHOOTING")
h1("困ったときは")
lead("たいていは、次のどれかです。")

TROUBLE = [
    ("「AI にアイデアを聞く」が押せない",
     [("ゲージが足りていないか、",), ("通信が切れています", True),
      ("。ノードを足すか、電波を確認してください。",)]),
    ("オフラインになった",
     [("画面の上に帯が出ます。",), ("編集はそのまま続けられます", True),
      ("。書いたものは端末に貯まり、つながった時点で自動的に同期されます。"
       "AI に聞く操作だけができません。",)]),
    ("マップが消えた",
     [("ログインしていない場合、マップは端末だけに保存され、",),
      ("最終更新から30日", True), ("で削除されます。ログインするとクラウドに残ります。",)]),
    ("AI の提案が的外れ",
     [("設定で ",), ("AI の人格", True),
      ("を変えられます。アドバイザー（肯定＋指摘）／上司（問いを返す）／"
       "アナリスト（論理と逆説）の3種類です。",)]),
    ("説明の言葉がむずかしい",
     [("説明のやさしさは",), ("誕生日から自動で決まります", True),
      ("。設定の誕生日が正しいか確認してください。",)]),
    ("画面がまぶしい・見づらい",
     [("ホーム画面の下、または設定で ",), ("ライト／ダーク", True),
      ("を切り替えられます。既定は端末の設定に従います。",)]),
    ("ログインのたびに数字を聞かれる",
     [("2要素認証が有効になっています。メールに届く",), ("6桁", True),
      ("を入れてください。一度通せば30日は聞かれません。設定で解除もできます。",)]),
]
tt = add_table(doc, len(TROUBLE) + 1, 2, [56, 114])
for i, htext in enumerate(["こうなった", "こうする"]):
    c = blank(tt.cell(0, i), (1, 2.5, 2.5, 2.5))
    cell_borders(c, LINE, sides=("bottom",))
    para(c, htext, size=6.8, bold=True, color=FAINT, spacing=1.3, tracking=1.2)
for j, (q, a) in enumerate(TROUBLE):
    cq = blank(tt.cell(j + 1, 0), (3.2, 2.5, 3.2, 2.5))
    ca = blank(tt.cell(j + 1, 1), (3.2, 2.5, 3.2, 2.5))
    for c in (cq, ca):
        cell_borders(c, HAIR, sides=("bottom",))
    para(cq, q, size=8.4, bold=True, spacing=1.75)
    rich(ca, a, size=8.4, spacing=1.75, base=MUTED)

spacer(doc, 12)
notebox(doc, "m-annoyed.png", "それでも直らないときは", [
    [("アプリ内の ",), ("「お問い合わせ」", True),
     (" から連絡してください。ホーム画面のいちばん下にあります。",)],
    [("送るときは、",), ("いつ・どの画面で・何をしたか", True),
     ("の3つが書いてあると、こちらで再現できます。",)],
])


# ---------------------------------------------------------------- 11. できること

new_page()
eyebrow("07 — INDEX")
h1("できること")
lead("この説明書で触れなかったものも含めた一覧です。")

INDEX_L = [("テーマを決める", "1枚に1つ"),
           ("ノードを足す", "選んで書いて追加"),
           ("AI にアイデアを聞く", "ゲージ1つ"),
           ("わからないを解く", "年齢に合わせて説明"),
           ("全体をレビュー", "根拠ノードが光る"),
           ("整える", "ツリー状に並べ直す"),
           ("完成にする", "バッジがもらえる")]
INDEX_R = [("共同編集", "招待リンクで同時に書く"),
           ("コミュニティ", "匿名で公開・コメント"),
           ("バッジ", "これまでの記録"),
           ("ライト／ダーク", "端末に合わせるが既定"),
           ("オフライン編集", "つながると自動で同期"),
           ("2要素認証", "メールに6桁"),
           ("アカウント削除", "設定から")]

idx = add_table(doc, len(INDEX_L), 4, [46, 39, 46, 39])
for j in range(len(INDEX_L)):
    for base, data in ((0, INDEX_L[j]), (2, INDEX_R[j])):
        ck = blank(idx.cell(j, base), (2.6, 0 if base == 0 else 5, 2.6, 1))
        cv = blank(idx.cell(j, base + 1), (2.6, 1, 2.6, 5 if base == 0 else 0))
        for c in (ck, cv):
            cell_borders(c, HAIR, sides=("bottom",))
        para(ck, data[0], size=8.4, bold=True, spacing=1.6)
        para(cv, data[1], size=7.8, color=MUTED, spacing=1.6, align=RIGHT)

picture(doc, asset("m-meditate.png"), 23, align=CENTER, before=38)
para(doc, "人間の脳で、考えよう。", size=9, bold=True, spacing=1.6, before=3,
     align=CENTER)
para(doc, "主役はいつも、あなたの思考です。", size=7.6, color=MUTED, spacing=1.8,
     align=CENTER)


fix_cells()

out = os.path.join(HERE, "思索Mindmap_取扱説明書.docx")
doc.save(out)
print("saved:", out)
